# -*- coding: utf-8 -*-
"""260611-29: PDF 병합 '2단 축소 배치' — 각 문서를 2-up(원본 용지 자동 회전)으로 줄여
배치하고 하단 쪽번호, 책갈피 페이지 재매핑, 목차·표지(Word 양식 우선·fitz 폴백) 생성.

설계 요약:
  - 원본 페이지가 세로(세로>가로)면 출력 시트=가로(좌우 2단), 가로면 출력 시트=세로(상하 2단).
    출력 시트 크기 = 원본 용지를 90° 회전한 크기(원본 용지 유지) → 2쪽을 최대 크기로.
  - 쪽번호 = 내용 시트(2단 시트) 번호(1-based). 목차·책갈피도 이 번호 기준.
  - 책갈피: 원본 page(1-based) → 내용 시트 = 파일시작시트 + (page-1)//2.
  - 최종 PDF = [표지] + [목차] + [2단 내용]. 책갈피 page 는 (표지+목차 수) 만큼 추가 오프셋.
"""
from __future__ import annotations

import os
import tempfile
from pathlib import Path

import fitz


class MergeCancelled(Exception):
    """진행 중 사용자가 취소."""


# 용지 크기(pt, 세로 기준)
PAGE_SIZES = {
    "A4": (595.28, 841.89), "Letter": (612.0, 792.0),
    "Legal": (612.0, 1008.0), "A3": (841.89, 1190.55),
}

DEFAULT_TWOUP = {
    "enabled": False,
    "nup": 2,                          # 260611-36: 한 장에 배치할 원본 쪽 수(2 또는 6)
    "page_size": "A4",                 # 출력 용지(PAGE_SIZES)
    "center": True,                    # 셀 안 가운데 정렬
    "fit_mode": "contain",             # 260611-46: 채움 방식 contain|cover|stretch
    "margin_top": 36, "margin_bottom": 48, "margin_left": 28, "margin_right": 28,  # pt
    "gap": 16,                         # 열(가로) 간격(pt)
    "gap_v": 16,                       # 행(세로) 간격(pt) — 6-up 등
    "crop_top": 0, "crop_bottom": 0, "crop_left": 0, "crop_right": 0,  # 원본 크롭(%)
    "duplex": False,                   # 양면(True)/단면(False) 인쇄
    "gutter": 0,                       # 제본용 여백(pt) — 단면=좌측, 양면=홀수 좌/짝수 우
    "facing_first": False,             # 260617-6: 맞쪽 인쇄 — 맨 앞 여백 페이지 1장 추가
    "doc_break": False,                # True면 문서마다 새 페이지(연속 채움 끔)
    "doc_start_odd": False,            # 260611-48: 양면+doc_break일 때 새 문서를 홀수 페이지에서 시작
    "border_outer": False, "border_h": False, "border_v": False,  # 외곽/내부 가로/세로 선
    "line_color": "#888888", "line_width": 1,                     # 선 색·굵기(pt)
    "margin_bg_on": False, "margin_bg": "#ffffff",  # 260611-44: 여백(시트 배경) 색
    "footer_pos": "center",           # left | center | right
    "footer_size": 11,
    # 260611-44: 쪽번호 블록·폰트
    "footer_block": False,            # 쪽번호 뒤 블록 그리기
    "footer_block_pad": 10,           # 글자보다 상하좌우 더 큰 비율(%) 0~20
    "footer_block_shape": "rect",     # rect | round (원형 테두리 직사각형)
    "footer_block_color": "#ffffff",  # 블록 색
    "footer_block_alpha": 100,        # 블록 투명도(%) 0~100
    "footer_font": "",                # 쪽번호 글꼴(빈값=기본 Helvetica)
    "footer_bold": False,             # 굵게
    "make_cover": True, "make_toc": True,
    "cover": {"title": "", "subtitle": "", "company": "", "name": ""},
    "cover_template": "",             # .docx (비우면 기본 생성)
    "toc_template": "",               # .docx (비우면 기본 생성)
    # 260611-51: 파일별 간지(divider)
    "make_divider": False,            # 각 파일 앞에 파일명 기반 간지 추가
    "divider_template": "",           # .docx (비우면 기본 생성)
    "divider_bg": "#eef2f7",          # 간지 배경색(여백색과 별도)
}


def merge_twoup_settings(s) -> dict:
    import copy
    out = copy.deepcopy(DEFAULT_TWOUP)
    if isinstance(s, dict):
        for k, v in s.items():
            if k == "cover" and isinstance(v, dict):
                out["cover"].update(v)
            else:
                out[k] = v
    return out


def _fit_rect(box, sw, sh):
    """source(sw×sh)를 box 안에 비율 유지·가운데 배치한 Rect."""
    if sw <= 0 or sh <= 0:
        return box
    scale = min(box.width / sw, box.height / sh)
    w = sw * scale; h = sh * scale
    x = box.x0 + (box.width - w) / 2.0
    y = box.y0 + (box.height - h) / 2.0
    return fitz.Rect(x, y, x + w, y + h)


def _hex_to_rgb(h):
    try:
        h = str(h).lstrip("#")
        return (int(h[0:2], 16) / 255.0, int(h[2:4], 16) / 255.0, int(h[4:6], 16) / 255.0)
    except Exception:
        return (0.5, 0.5, 0.5)


def _draw_grid_lines(page, boxes, s):
    """260611-43: 외곽선·내부 가로/세로선(선색·굵기). boxes=fitz.Rect 목록(격자)."""
    if not boxes:
        return
    outer = bool(s.get("border_outer", False))
    bh = bool(s.get("border_h", False))
    bv = bool(s.get("border_v", False))
    if not (outer or bh or bv):
        return
    col = _hex_to_rgb(s.get("line_color", "#888888"))
    w = max(0.2, float(s.get("line_width", 1)))
    x0 = min(b.x0 for b in boxes); y0 = min(b.y0 for b in boxes)
    x1 = max(b.x1 for b in boxes); y1 = max(b.y1 for b in boxes)
    cols = max(1, sum(1 for b in boxes if abs(b.y0 - boxes[0].y0) < 1.0))
    rows = max(1, len(boxes) // cols)
    try:
        if outer:
            page.draw_rect(fitz.Rect(x0, y0, x1, y1), color=col, width=w)
        if bv and cols > 1:
            for c in range(1, cols):
                xa = (boxes[c - 1].x1 + boxes[c].x0) / 2.0
                page.draw_line(fitz.Point(xa, y0), fitz.Point(xa, y1), color=col, width=w)
        if bh and rows > 1:
            for r in range(1, rows):
                ya = (boxes[(r - 1) * cols].y1 + boxes[r * cols].y0) / 2.0
                page.draw_line(fitz.Point(x0, ya), fitz.Point(x1, ya), color=col, width=w)
    except Exception:
        pass


_FONT_CACHE = {}


def _win_font_file(family, bold=False):
    """Windows 글꼴 레지스트리에서 family(+굵기)에 맞는 폰트 파일 경로. 못 찾으면 None."""
    if not family:
        return None
    key = (str(family).lower(), bool(bold))
    if key in _FONT_CACHE:
        return _FONT_CACHE[key]
    path = None
    try:
        import os, winreg
        fam = str(family).lower()
        reg = winreg.OpenKey(winreg.HKEY_LOCAL_MACHINE,
                             r"SOFTWARE\Microsoft\Windows NT\CurrentVersion\Fonts")
        fdir = os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts")
        cand_exact = None; cand_any = None
        i = 0
        while True:
            try:
                name, val, _ = winreg.EnumValue(reg, i)
            except OSError:
                break
            i += 1
            nl = name.lower()
            if fam in nl:
                p = val if os.path.isabs(val) else os.path.join(fdir, val)
                is_bold = "bold" in nl
                if is_bold == bool(bold) and cand_exact is None:
                    cand_exact = p
                if cand_any is None:
                    cand_any = p
        path = cand_exact or cand_any
    except Exception:
        path = None
    _FONT_CACHE[key] = path
    return path


def _draw_sheet_bg(page, s):
    """260611-44: 여백 색 — 시트 배경 전체를 채움(내용 쪽은 위에 그려져 여백만 보임)."""
    if not bool(s.get("margin_bg_on", False)):
        return
    try:
        page.draw_rect(page.rect, color=None, fill=_hex_to_rgb(s.get("margin_bg", "#ffffff")))
    except Exception:
        pass


def _draw_footer(page, text, s, ow):
    pos = s.get("footer_pos", "center")
    if pos == "none":                       # 260617-6: 쪽번호 표시 안 함
        return
    fs = int(s.get("footer_size", 11))
    mb = float(s.get("margin_bottom", 48))
    mt = float(s.get("margin_top", 36))
    ml = float(s.get("margin_left", 28)); mr = float(s.get("margin_right", 28))
    y = page.rect.height - mb / 2.0
    bold = bool(s.get("footer_bold", False))
    fam = s.get("footer_font", "")
    fontfile = _win_font_file(fam, bold) if fam else None
    builtin = "hebo" if (bold and not fontfile) else "helv"   # 굵게(커스텀 글꼴 없을 때)
    tw = fitz.get_text_length(text, fontname=builtin, fontsize=fs)
    if pos == "left":
        x = ml
    elif pos == "right":
        x = ow - mr - tw
    elif pos == "topright":                 # 260617-6: 우상단
        x = ow - mr - tw
        y = max(fs + 2.0, mt * 0.7)
    else:
        x = (ow - tw) / 2.0
    # 260611-44: 쪽번호 블록(글자보다 0~20% 크게)
    if bool(s.get("footer_block", False)):
        pad = max(0, min(20, int(s.get("footer_block_pad", 10)))) / 100.0
        padx = max(2.0, tw * pad); pady = max(2.0, fs * pad)
        rect = fitz.Rect(x - padx, y - fs * 0.82 - pady, x + tw + padx, y + fs * 0.22 + pady)
        col = _hex_to_rgb(s.get("footer_block_color", "#ffffff"))
        alpha = max(0, min(100, int(s.get("footer_block_alpha", 100)))) / 100.0
        kw = dict(color=None, fill=col, fill_opacity=alpha, width=0)
        try:
            if s.get("footer_block_shape", "rect") == "round":
                page.draw_rect(rect, radius=0.4, **kw)
            else:
                page.draw_rect(rect, **kw)
        except Exception:
            try:
                page.draw_rect(rect, **kw)
            except Exception:
                pass
    try:
        if fontfile:
            page.insert_text((x, y), text, fontfile=fontfile, fontname="pnf",
                             fontsize=fs, color=(0.1, 0.1, 0.1))
        else:
            page.insert_text((x, y), text, fontname=builtin, fontsize=fs, color=(0.1, 0.1, 0.1))
    except Exception:
        page.insert_text((x, y), text, fontname="helv", fontsize=fs, color=(0.1, 0.1, 0.1))


def _paper(s):
    """선택 용지의 (short, long) pt."""
    w, h = PAGE_SIZES.get(s.get("page_size", "A4"), PAGE_SIZES["A4"])
    return (min(w, h), max(w, h))


def _grid_dims(p0, s):
    """nup·원본 방향으로 (cols, rows, 용지 가로, 용지 세로) 결정.
    2-up: 세로 원본→가로 용지 2×1 / 가로 원본→세로 용지 1×2.
    6-up: 세로 용지 고정 2열×3행."""
    short, long_ = _paper(s)
    nup = int(s.get("nup", 2) or 2)
    portrait_src = p0.height >= p0.width
    if nup >= 8:
        return 2, 4, short, long_                 # 세로 용지 2×4
    if nup >= 6:
        return 2, 3, short, long_                 # 세로 용지 2×3
    if nup >= 4:
        return 2, 2, short, long_                 # 260617-6: 세로 용지 2×2(4-up)
    # 2-up
    if portrait_src:
        return 2, 1, long_, short                 # 가로 용지, 좌우
    return 1, 2, short, long_                      # 세로 용지, 상하


def _grid_layout(p0, s, sheet_no=1):
    """반환: (ow, oh, [cell rects]) — cols×rows 격자, 열=gap·행=gap_v 간격, 가운데 정렬.
    제본 여백(gutter): 단면=좌측, 양면=홀수 시트 좌·짝수 시트 우에 추가."""
    cols, rows, ow, oh = _grid_dims(p0, s)
    mt = float(s.get("margin_top", 36)); mb = float(s.get("margin_bottom", 48))
    ml = float(s.get("margin_left", 28)); mr = float(s.get("margin_right", 28))
    gutter = max(0.0, float(s.get("gutter", 0)))
    if gutter:
        if bool(s.get("duplex", False)) and int(sheet_no) % 2 == 0:
            mr += gutter                 # 양면 짝수 시트 → 우측 제본 여백
        else:
            ml += gutter                 # 단면 또는 양면 홀수 시트 → 좌측
    gh = float(s.get("gap", 16)); gv = float(s.get("gap_v", 16))
    footer_h = float(s.get("footer_size", 11)) + 10.0
    avail_w = ow - ml - mr - gh * (cols - 1)
    avail_h = oh - mt - mb - footer_h - gv * (rows - 1)
    cw = avail_w / cols
    ch = avail_h / rows
    boxes = []
    for r in range(rows):              # 위→아래, 좌→우 읽기 순서
        for c in range(cols):
            x0 = ml + c * (cw + gh)
            y0 = mt + r * (ch + gv)
            boxes.append(fitz.Rect(x0, y0, x0 + cw, y0 + ch))
    return ow, oh, boxes


def _crop_clip(src_rect, s):
    """원본 rect에서 상하좌우 %를 잘라낸 clip rect(가운데 영역). 0이면 전체."""
    ct = max(0.0, min(45.0, float(s.get("crop_top", 0)))) / 100.0
    cb = max(0.0, min(45.0, float(s.get("crop_bottom", 0)))) / 100.0
    cl = max(0.0, min(45.0, float(s.get("crop_left", 0)))) / 100.0
    cr = max(0.0, min(45.0, float(s.get("crop_right", 0)))) / 100.0
    w, h = src_rect.width, src_rect.height
    return fitz.Rect(src_rect.x0 + w * cl, src_rect.y0 + h * ct,
                     src_rect.x1 - w * cr, src_rect.y1 - h * cb)


def _cover_clip(clip, box):
    """260611-46: box(셀) 가로세로비에 꽉 차도록 clip을 가운데에서 더 잘라낸 sub-rect."""
    if box.width <= 0 or box.height <= 0 or clip.width <= 0 or clip.height <= 0:
        return clip
    tar = box.width / box.height
    src = clip.width / clip.height
    if src > tar:                       # 원본이 더 넓음 → 좌우 잘라냄
        nw = clip.height * tar
        dx = (clip.width - nw) / 2.0
        return fitz.Rect(clip.x0 + dx, clip.y0, clip.x1 - dx, clip.y1)
    nh = clip.width / tar               # 원본이 더 높음 → 상하 잘라냄
    dy = (clip.height - nh) / 2.0
    return fitz.Rect(clip.x0, clip.y0 + dy, clip.x1, clip.y1 - dy)


def _place_page(page, src, pno, box, s):
    """260611-46: 채움 방식(fit_mode)에 따라 원본 쪽을 셀(box)에 삽입.
    contain=비율 유지·여백 / cover=비율 유지·가장자리 잘라 꽉 채움 / stretch=비율 무시."""
    clip = _crop_clip(src[pno].rect, s)
    mode = s.get("fit_mode", "contain")
    if mode == "cover":
        tgt = box
    elif mode == "stretch":
        tgt = box
    else:
        tgt = _fit_rect(box, clip.width, clip.height)
    # 260611-48: 여백 색이 켜진 경우, 문서가 놓이는 영역엔 흰 종이를 깔아
    #            문서 내부 투명 영역에 여백색이 비쳐 보이지 않게 함.
    if bool(s.get("margin_bg_on", False)):
        try:
            page.draw_rect(tgt, color=None, fill=(1, 1, 1))
        except Exception:
            pass
    try:
        if mode == "cover":
            page.show_pdf_page(box, src, pno, clip=_cover_clip(clip, box))
        elif mode == "stretch":
            # keep_proportion=False 라야 셀에 꽉 차게 비율을 늘림(기본 True면 비율 유지·가운데)
            page.show_pdf_page(box, src, pno, clip=clip, keep_proportion=False)
        else:
            page.show_pdf_page(tgt, src, pno, clip=clip)
    except Exception:
        try:
            page.show_pdf_page(_fit_rect(box, clip.width, clip.height), src, pno)
        except Exception:
            pass


def _impose_item(out_doc, src, s, content_sheet, on_sheet=None):
    """src(fitz 문서)를 2-up 으로 out_doc 에 추가. 방향은 '짝(시트)별'로 결정해
    혼합 방향 문서도 페이지마다 올바르게 배치. on_sheet(): 시트마다 호출(진행/취소).
    반환: (시트 수, 시작 내용시트번호)."""
    n = src.page_count
    if n == 0:
        return 0, content_sheet
    nup = max(1, int(s.get("nup", 2) or 2))
    start = content_sheet
    sheets = 0
    for i in range(0, n, nup):
        ow, oh, boxes = _grid_layout(src[i].rect, s)    # 이 짝의 방향·격자
        page = out_doc.new_page(width=ow, height=oh)
        for j, box in enumerate(boxes):
            pno = i + j
            if pno >= n:
                break
            _place_page(page, src, pno, box, s)          # 크롭 + 채움 방식
        _draw_footer(page, str(content_sheet), s, ow)
        content_sheet += 1
        sheets += 1
        if on_sheet is not None:
            on_sheet()
    return sheets, start


def _open_item_doc(it):
    """항목(pdf/shots)을 fitz 문서로. shots 는 이미지 페이지 임시 문서."""
    if isinstance(it, str):
        return fitz.open(it)
    if it.get("type") == "shots":
        d = fitz.open()
        for img in (it.get("paths") or []):
            try:
                pix = fitz.Pixmap(img)
                pg = d.new_page(width=pix.width, height=pix.height)
                pg.insert_image(fitz.Rect(0, 0, pix.width, pix.height), filename=img)
            except Exception:
                continue
        return d
    return fitz.open(str(it["path"]))


# 260611-49/50: 미리보기 전용 캐시 — 같은 샘플을 매 렌더마다 다시 열지 않게.
#   _PREVIEW_PC: 쪽 수 캐시(가벼움) — 펼침 계산에만 사용(문서 핸들 보관 안 함).
#   _PREVIEW_CACHE: 열린 문서 캐시(LRU) — 현재 시트 렌더용. 'pinned' 키는 절대 닫지 않음
#   (사용 중 문서를 eviction이 닫아버리는 use-after-close 크래시 방지).
_PREVIEW_CACHE = {}
_PREVIEW_ORDER = []
_PREVIEW_CAP = 12
_PREVIEW_PINNED = set()
_PREVIEW_PC = {}


def _preview_key(it):
    try:
        if isinstance(it, dict) and it.get("type") == "shots":
            return ("shots", tuple(it.get("paths") or []))
        p = it if isinstance(it, str) else str(it.get("path", ""))
        if not p:
            return None
        stt = os.stat(p)
        return ("file", p, int(stt.st_mtime), int(stt.st_size))
    except Exception:
        return None


def _evict_preview():
    """cap 초과분 정리 — pinned(사용 중) 문서는 건너뜀."""
    i = 0
    while len(_PREVIEW_ORDER) > _PREVIEW_CAP and i < len(_PREVIEW_ORDER):
        key = _PREVIEW_ORDER[i]
        if key in _PREVIEW_PINNED:
            i += 1
            continue
        _PREVIEW_ORDER.pop(i)
        od = _PREVIEW_CACHE.pop(key, None)
        try:
            if od is not None:
                od.close()
        except Exception:
            pass


def _open_item_cached(it):
    """반환 (doc, cached). cached=True면 캐시 소유이므로 호출자가 닫지 말 것.
    캐시된 문서는 pinned 처리 — 반드시 _unpin_preview()로 해제해야 eviction 대상이 됨."""
    key = _preview_key(it)
    if key is None:
        return _open_item_doc(it), False
    d = _PREVIEW_CACHE.get(key)
    if d is not None:
        try:
            _ = d.page_count            # 닫힌 문서 접근 시 예외 → 캐시 무효화
            _PREVIEW_ORDER.remove(key); _PREVIEW_ORDER.append(key)
            _PREVIEW_PINNED.add(key)
            return d, True
        except Exception:
            _PREVIEW_CACHE.pop(key, None)
            try:
                _PREVIEW_ORDER.remove(key)
            except ValueError:
                pass
    d = _open_item_doc(it)
    _PREVIEW_CACHE[key] = d; _PREVIEW_ORDER.append(key); _PREVIEW_PINNED.add(key)
    _evict_preview()
    return d, True


def _unpin_preview():
    """현재 렌더가 끝나 사용 중 표시 해제 + cap 정리."""
    _PREVIEW_PINNED.clear()
    _evict_preview()


def _item_pagecount(it):
    """쪽 수(캐시). 문서 핸들을 보관하지 않으므로 항목이 많아도 안전."""
    key = _preview_key(it)
    if key is not None and key in _PREVIEW_PC:
        return _PREVIEW_PC[key]
    try:
        d = _open_item_doc(it)
        try:
            pc = d.page_count
        finally:
            d.close()
    except Exception:
        pc = 0
    if key is not None:
        _PREVIEW_PC[key] = pc
    return pc


# 미리보기 계획 캐시 — 설정/샘플이 바뀔 때만 재계획. 페이지는 요청 시 1장만 렌더.
_PREVIEW_DOC = {"sig": None}


def _close_preview_doc():
    st = _PREVIEW_DOC
    for key in ("cover_doc", "toc_doc", "fallback"):
        d = st.get(key)
        if d is not None:
            try:
                d.close()
            except Exception:
                pass
    for d in (st.get("div_docs") or {}).values():
        try:
            d.close()
        except Exception:
            pass
    st.clear(); st["sig"] = None


def _preview_sig(items, s):
    import json
    keys = tuple(_preview_key(it) for it in items)
    try:
        ss = json.dumps(s, sort_keys=True, default=str, ensure_ascii=False)
    except Exception:
        ss = repr(sorted(((str(k), str(v)) for k, v in s.items())))
    return (keys, ss)


def clear_preview_cache():
    """미리보기 창을 닫을 때 캐시된 문서 핸들 정리."""
    for d in list(_PREVIEW_CACHE.values()):
        try:
            d.close()
        except Exception:
            pass
    _PREVIEW_CACHE.clear(); _PREVIEW_ORDER.clear()
    _PREVIEW_PINNED.clear(); _PREVIEW_PC.clear()
    _close_preview_doc()


def _build_preview_state(items, s):
    """미리보기 계획(슬랩·번호·표지/목차 문서·크기)만 만들어 캐시. 내용 시트는 렌더 안 함."""
    st = _PREVIEW_DOC
    slabs, file_blocks, _p2s = _merge_slabs(items, s, _item_pagecount)
    if not slabs:
        raise RuntimeError("배치할 내용이 없습니다.")
    fw = fh = None
    for sl in slabs:
        if sl["kind"] == "content":
            for cell in sl["chunk"]:
                if cell:
                    it, p = cell
                    d, _c = _open_item_cached(it)
                    try:
                        fw, fh, _b = _grid_layout(d[p].rect, s)
                    except Exception:
                        fw = None
                    break
            if fw:
                break
    if fw is None:
        shp, lop = _paper(s); fw, fh, _b = _grid_layout(fitz.Rect(0, 0, shp, lop), s)
    make_cover = bool(s.get("make_cover", True)); make_toc = bool(s.get("make_toc", True))
    cover_doc = _fitz_cover_doc(s.get("cover", {}), fw, fh) if make_cover else None
    cover_pages = cover_doc.page_count if cover_doc else 0
    names = [fb["name"] for fb in file_blocks]
    toc_doc = None; toc_pages = 0
    if make_toc:
        toc_doc = _fitz_toc_doc([(n, "") for n in names], fw, fh)
        toc_pages = toc_doc.page_count
    front_pages = cover_pages + toc_pages
    body_plan, _slab_phys, toc_entries = _merge_body_plan(slabs, file_blocks, s, front_pages)
    if make_toc:
        try:
            toc_doc.close()
        except Exception:
            pass
        toc_doc = _fitz_toc_doc(toc_entries, fw, fh)
        for r in range(toc_doc.page_count):
            _draw_footer(toc_doc[r], _roman(r + 1), s, toc_doc[r].rect.width)
    st.update({"cover_doc": cover_doc, "toc_doc": toc_doc, "cover_pages": cover_pages,
               "front_pages": front_pages, "body_plan": body_plan, "slabs": slabs,
               "fw": fw, "fh": fh, "s": s, "div_docs": {}, "fallback": None,
               "total": front_pages + len(body_plan)})
    _unpin_preview()


def _render_preview_page(idx, dpi):
    st = _PREVIEW_DOC
    if st.get("fallback") is not None:
        pg = st["fallback"][0]
    else:
        cover_pages = st["cover_pages"]; front_pages = st["front_pages"]; s = st["s"]
        cells = []; temp = None
        if idx < cover_pages:
            pg = st["cover_doc"][idx]
        elif idx < front_pages:
            pg = st["toc_doc"][idx - cover_pages]
        else:
            d = st["body_plan"][idx - front_pages]; fw, fh = st["fw"], st["fh"]
            if d["kind"] == "blank":
                temp = fitz.open(); pg = temp.new_page(width=fw, height=fh); _draw_sheet_bg(pg, s)
            elif d["kind"] == "divider":
                sl = st["slabs"][d["slab"]]; fi = sl["file"]
                dd = st["div_docs"].get(fi)
                if dd is None:
                    dd = _fitz_divider_doc(sl["name"], fw, fh, s.get("divider_bg", "#eef2f7"))
                    st["div_docs"][fi] = dd
                pg = dd[0]
            else:
                opened = {}

                def _doc(it):
                    k = id(it)
                    if k not in opened:
                        opened[k] = _open_item_cached(it)
                    return opened[k][0]
                temp = fitz.open()
                cells = _render_content_sheet(temp, st["slabs"][d["slab"]]["chunk"], s,
                                              d["phys"], _doc)
                pg = temp[0]
                if d["num"]:
                    _draw_footer(pg, d["num"], s, pg.rect.width)
                for dch, cached in opened.values():
                    if not cached:
                        try:
                            dch.close()
                        except Exception:
                            pass
                _unpin_preview()
        ow, oh = pg.rect.width, pg.rect.height
        pix = pg.get_pixmap(matrix=fitz.Matrix(dpi / 72.0, dpi / 72.0))
        png = pix.tobytes("png")
        if temp is not None:
            try:
                temp.close()
            except Exception:
                pass
        return png, ow, oh, cells
    ow, oh = pg.rect.width, pg.rect.height
    pix = pg.get_pixmap(matrix=fitz.Matrix(dpi / 72.0, dpi / 72.0))
    return pix.tobytes("png"), ow, oh, []


def compose_preview(items_or_path, settings, sheet_index=0, dpi=110):
    """미리보기용 — 표지·목차·간지·본문을 포함하되 '요청한 페이지 1장만' 렌더.
    설정/샘플이 바뀔 때만 계획을 다시 만들고, 페이지 넘김은 계획 재사용.
    반환: (png_bytes, ow, oh, [cell rects], total_pages)."""
    s = merge_twoup_settings(settings)
    items = ([{"type": "pdf", "path": items_or_path}]
             if isinstance(items_or_path, str) else list(items_or_path or []))
    sig = _preview_sig(items, s)
    if _PREVIEW_DOC.get("sig") != sig:
        _close_preview_doc()
        try:
            _build_preview_state(items, s)
        except Exception:
            _close_preview_doc()
            f = fitz.open(); shp, lop = _paper(s); f.new_page(width=shp, height=lop)
            _PREVIEW_DOC["fallback"] = f; _PREVIEW_DOC["total"] = 1
        _PREVIEW_DOC["sig"] = sig
    total = max(1, _PREVIEW_DOC.get("total", 1))
    idx = max(0, min(total - 1, int(sheet_index)))
    png, ow, oh, cells = _render_preview_page(idx, dpi)
    return png, ow, oh, cells, total


def _docx_to_pdf(docx_path, pdf_path):
    """Word 자동화로 .docx→PDF. 성공 시 True."""
    try:
        import pythoncom
        import win32com.client
    except Exception:
        return False
    word = None
    try:
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(str(Path(docx_path).resolve()))
        doc.SaveAs(str(Path(pdf_path).resolve()), FileFormat=17)  # wdFormatPDF
        doc.Close(False)
        return os.path.exists(pdf_path)
    except Exception:
        return False
    finally:
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


def _fill_cover_docx(template, cover, out_docx):
    """표지 양식 docx 의 {{TITLE}}/{{SUBTITLE}}/{{COMPANY}}/{{NAME}} 치환(없으면 기본 생성)."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        return False
    repl = {"{{TITLE}}": cover.get("title", ""), "{{SUBTITLE}}": cover.get("subtitle", ""),
            "{{COMPANY}}": cover.get("company", ""), "{{NAME}}": cover.get("name", "")}
    try:
        if template and os.path.exists(template):
            doc = Document(template)
            for p in doc.paragraphs:
                for k, v in repl.items():
                    if k in p.text:
                        for r in p.runs:
                            if k in r.text:
                                r.text = r.text.replace(k, v)
                        # 런 분할로 못 바꾼 경우 단락 텍스트 전체 치환
                        if k in p.text:
                            p.text = p.text.replace(k, v)
        else:
            doc = Document()
            for _ in range(6):
                doc.add_paragraph("")
            h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.add_run(cover.get("title", "")); run.bold = True; run.font.size = Pt(32)
            sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = sub.add_run(cover.get("subtitle", "")); sr.font.size = Pt(18)
            for _ in range(10):
                doc.add_paragraph("")
            for key in ("company", "name"):
                pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                rr = pp.add_run(cover.get(key, "")); rr.font.size = Pt(14)
        doc.save(out_docx)
        return True
    except Exception:
        return False


def _fill_toc_docx(template, entries, out_docx):
    """목차 양식 docx 의 {{TOC}} 단락을 '파일명 ... 쪽' 목록으로 치환(없으면 기본 생성)."""
    try:
        from docx import Document
        from docx.shared import Pt
    except Exception:
        return False
    try:
        if template and os.path.exists(template):
            doc = Document(template)
            anchor = None
            for p in doc.paragraphs:
                if "{{TOC}}" in p.text:
                    anchor = p
                    break
            if anchor is not None:
                anchor.text = ""
                for name, pg in entries:
                    anchor.insert_paragraph_before(f"{name} ............ {pg}")
            else:
                for name, pg in entries:
                    doc.add_paragraph(f"{name} ............ {pg}")
        else:
            doc = Document()
            h = doc.add_paragraph(); r = h.add_run("목차"); r.bold = True; r.font.size = Pt(24)
            doc.add_paragraph("")
            for name, pg in entries:
                doc.add_paragraph(f"{name} ............ {pg}")
        doc.save(out_docx)
        return True
    except Exception:
        return False


def _fitz_cover_pdf(cover, ow, oh, out_pdf):
    d = fitz.open()
    pg = d.new_page(width=ow, height=oh)
    cx = ow / 2.0
    pg.insert_text((cx - fitz.get_text_length(cover.get("title", ""), fontsize=28) / 2,
                    oh * 0.32), cover.get("title", ""), fontsize=28, color=(0, 0, 0))
    pg.insert_text((cx - fitz.get_text_length(cover.get("subtitle", ""), fontsize=16) / 2,
                    oh * 0.32 + 40), cover.get("subtitle", ""), fontsize=16, color=(0.2, 0.2, 0.2))
    pg.insert_text((cx - fitz.get_text_length(cover.get("company", ""), fontsize=14) / 2,
                    oh * 0.72), cover.get("company", ""), fontsize=14, color=(0.1, 0.1, 0.1))
    pg.insert_text((cx - fitz.get_text_length(cover.get("name", ""), fontsize=14) / 2,
                    oh * 0.72 + 24), cover.get("name", ""), fontsize=14, color=(0.1, 0.1, 0.1))
    d.save(out_pdf); d.close()
    return True


def _fitz_toc_pdf(entries, ow, oh, out_pdf):
    d = fitz.open()
    pg = d.new_page(width=ow, height=oh)
    pg.insert_text((50, 60), "목차", fontsize=22, color=(0, 0, 0))
    y = 110
    for name, p in entries:
        if y > oh - 50:
            pg = d.new_page(width=ow, height=oh); y = 60
        line = f"{name}"
        pg.insert_text((50, y), line, fontsize=12, color=(0, 0, 0))
        ptxt = str(p)
        pg.insert_text((ow - 50 - fitz.get_text_length(ptxt, fontsize=12), y),
                       ptxt, fontsize=12, color=(0, 0, 0))
        y += 26
    d.save(out_pdf); d.close()
    return True


# ── 260611-51: 표지(번호없음)·목차(로마자)·간지(파일별)·본문(아라비아) 통합 조립 ──

def _roman(n):
    """1→i, 2→ii … 소문자 로마 숫자."""
    n = int(n)
    if n <= 0:
        return str(n)
    table = [(1000, "m"), (900, "cm"), (500, "d"), (400, "cd"), (100, "c"),
             (90, "xc"), (50, "l"), (40, "xl"), (10, "x"), (9, "ix"),
             (5, "v"), (4, "iv"), (1, "i")]
    out = ""
    for v, sym in table:
        while n >= v:
            out += sym; n -= v
    return out


_KR_FONT = {}      # bold -> (fontfile|None, fitz.Font|None) — 반복 로드 방지(성능)


def _kr_font(bold=False):
    k = bool(bold)
    if k not in _KR_FONT:
        ff = _win_font_file("맑은 고딕", k) or _win_font_file("malgun", k)
        fo = None
        if ff:
            try:
                fo = fitz.Font(fontfile=ff)
            except Exception:
                fo = None
        _KR_FONT[k] = (ff, fo)
    return _KR_FONT[k]


def _kr_text(page, pt, text, fontsize, color=(0, 0, 0), bold=False):
    """한글 가능한 텍스트 삽입(맑은 고딕 → 실패 시 helv)."""
    ff, _fo = _kr_font(bold)
    try:
        if ff:
            page.insert_text(pt, text, fontfile=ff, fontname="krf", fontsize=fontsize, color=color)
            return
    except Exception:
        pass
    page.insert_text(pt, text, fontname=("hebo" if bold else "helv"), fontsize=fontsize, color=color)


def _kr_width(text, fontsize, bold=False):
    _ff, fo = _kr_font(bold)
    if fo is not None:
        try:
            return fo.text_length(text, fontsize)
        except Exception:
            pass
    return fitz.get_text_length(text, fontname=("hebo" if bold else "helv"), fontsize=fontsize)


def _fitz_cover_doc(cover, w, h):
    d = fitz.open(); pg = d.new_page(width=w, height=h)
    cx = w / 2.0
    rows = [(cover.get("title", ""), 30, h * 0.30, True),
            (cover.get("subtitle", ""), 17, h * 0.30 + 44, False),
            (cover.get("company", ""), 14, h * 0.72, False),
            (cover.get("name", ""), 14, h * 0.72 + 24, False)]
    for txt, fs, y, bold in rows:
        if not txt:
            continue
        _kr_text(pg, (cx - _kr_width(txt, fs, bold) / 2.0, y), txt, fs, (0.05, 0.05, 0.05), bold)
    return d


def _fitz_toc_doc(entries, w, h):
    d = fitz.open(); pg = d.new_page(width=w, height=h)
    _kr_text(pg, (50, 64), "목차", 22, (0, 0, 0), True)
    y = 112
    for name, pnum in entries:
        if y > h - 50:
            pg = d.new_page(width=w, height=h); y = 64
        _kr_text(pg, (50, y), str(name), 12, (0, 0, 0))
        ptxt = str(pnum)
        _kr_text(pg, (w - 50 - _kr_width(ptxt, 12), y), ptxt, 12, (0, 0, 0))
        y += 26
    return d


def _fill_divider_docx(template, title, out_docx):
    """간지 양식의 {{TITLE}}/{{NAME}} 을 파일명으로 치환(없으면 기본 생성)."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        return False
    try:
        if template and os.path.exists(template):
            doc = Document(template)
            for p in doc.paragraphs:
                for k in ("{{TITLE}}", "{{NAME}}"):
                    if k in p.text:
                        for r in p.runs:
                            if k in r.text:
                                r.text = r.text.replace(k, title)
                        if k in p.text:
                            p.text = p.text.replace(k, title)
        else:
            doc = Document()
            for _ in range(8):
                doc.add_paragraph("")
            h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = h.add_run(title); run.bold = True; run.font.size = Pt(30)
        doc.save(out_docx)
        return True
    except Exception:
        return False


def _fitz_divider_doc(title, w, h, bg):
    d = fitz.open(); pg = d.new_page(width=w, height=h)
    try:
        pg.draw_rect(pg.rect, color=None, fill=_hex_to_rgb(bg))
    except Exception:
        pass
    fs = 30
    while fs > 12 and _kr_width(title, fs, True) > w * 0.82:
        fs -= 2
    _kr_text(pg, ((w - _kr_width(title, fs, True)) / 2.0, h * 0.5), title, fs, (0.1, 0.1, 0.1), True)
    return d


def _paint_bg_behind(doc, bg):
    """Word 등으로 만든 문서 각 페이지 뒤에 배경색을 깔기(내용 아래)."""
    try:
        col = _hex_to_rgb(bg)
        for pg in doc:
            pg.draw_rect(pg.rect, color=None, fill=col, overlay=False)
    except Exception:
        pass


def write_sample_templates(folder):
    """표지·목차·간지 Word 양식 샘플을 folder 에 생성. 생성된 경로 목록 반환."""
    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except Exception:
        return []
    made = []
    folder = Path(folder)

    def _save(name, build):
        try:
            doc = Document(); build(doc)
            p = str(folder / name); doc.save(p); made.append(p)
        except Exception:
            pass

    def cover(doc):
        for _ in range(6):
            doc.add_paragraph("")
        for ph, fs, bold in [("{{TITLE}}", 32, True), ("{{SUBTITLE}}", 18, False)]:
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(ph); r.bold = bold; r.font.size = Pt(fs)
        for _ in range(10):
            doc.add_paragraph("")
        for ph in ("{{COMPANY}}", "{{NAME}}"):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(ph).font.size = Pt(14)

    def toc(doc):
        h = doc.add_paragraph(); r = h.add_run("목차"); r.bold = True; r.font.size = Pt(24)
        doc.add_paragraph("")
        p = doc.add_paragraph(); p.add_run("{{TOC}}")     # 이 단락이 항목 목록으로 치환됨

    def divider(doc):
        for _ in range(8):
            doc.add_paragraph("")
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("{{TITLE}}"); r.bold = True; r.font.size = Pt(30)

    _save("표지_샘플.docx", cover)
    _save("목차_샘플.docx", toc)
    _save("간지_샘플.docx", divider)
    return made


def _item_name(it):
    if isinstance(it, dict):
        return it.get("name") or Path(str(it.get("path", "항목"))).stem
    return Path(str(it)).stem


def _merge_slabs(items, s, pc_fn):
    """본문 슬랩(간지/내용) 구성. 렌더링 없이 쪽 수만 사용(가벼움).
    반환: (slabs, file_blocks, page_to_slab)."""
    nup = max(1, int(s.get("nup", 2) or 2))
    make_divider = bool(s.get("make_divider", False))
    doc_break = bool(s.get("doc_break", False))
    per_file = make_divider or doc_break
    live = [(i, it) for i, it in enumerate(items) if pc_fn(it) > 0]
    slabs = []; page_to_slab = {}; file_blocks = []
    if per_file:
        for i, it in live:
            pc = pc_fn(it); lo = len(slabs)
            if make_divider:
                slabs.append({"kind": "divider", "file": i, "name": _item_name(it)})
            for k in range(0, pc, nup):
                chunk = [(it, p) for p in range(k, min(k + nup, pc))]
                for (cit, cp) in chunk:
                    page_to_slab[(id(cit), cp)] = len(slabs)
                slabs.append({"kind": "content", "file": i, "name": _item_name(it), "chunk": chunk})
            file_blocks.append({"file": i, "it": it, "name": _item_name(it),
                                "slab_lo": lo, "slab_hi": len(slabs)})
    else:
        flat = [(it, p, i) for (i, it) in live for p in range(pc_fn(it))]
        for k in range(0, len(flat), nup):
            grp = flat[k:k + nup]
            chunk = [(it, p) for (it, p, _i) in grp]
            for (it, p, _i) in grp:
                page_to_slab[(id(it), p)] = len(slabs)
            slabs.append({"kind": "content", "file": grp[0][2], "name": _item_name(grp[0][0]),
                          "chunk": chunk})
        seen = set()
        for pos, (it, p, i) in enumerate(flat):
            if i not in seen:
                seen.add(i)
                file_blocks.append({"file": i, "it": it, "name": _item_name(it),
                                    "slab_lo": pos // nup, "slab_hi": None})
    return slabs, file_blocks, page_to_slab


def _merge_body_plan(slabs, file_blocks, s, front_pages):
    """본문 페이지 순서/번호/물리위치 계획(렌더링 없음).
    반환: (body_plan[{kind,slab,num,phys}], slab_phys, toc_entries[(name,arabic)])."""
    nup = max(1, int(s.get("nup", 2) or 2))
    make_divider = bool(s.get("make_divider", False))
    doc_break = bool(s.get("doc_break", False))
    duplex = bool(s.get("duplex", False))
    per_file = make_divider or doc_break
    odd = bool(s.get("doc_start_odd", False)) and duplex and per_file
    body_plan = []; slab_phys = {}
    cur = [front_pages]; ara = [0]

    def emit(kind, slab=None):
        cur[0] += 1
        if kind in ("divider", "blank", "content"):
            ara[0] += 1
        d = {"kind": kind, "slab": slab, "num": str(ara[0]) if kind == "content" else None,
             "phys": cur[0]}
        body_plan.append(d)
        if slab is not None:
            slab_phys[slab] = cur[0]
        return d

    if per_file:
        for fb in file_blocks:
            lo, hi = fb["slab_lo"], fb["slab_hi"]
            has_div = slabs[lo]["kind"] == "divider"
            if odd:
                if (cur[0] + 1) % 2 == 0:
                    emit("blank")
                if has_div:
                    emit("divider", slab=lo)
                    if (cur[0] + 1) % 2 == 0:
                        emit("blank")
                    for si in range(lo + 1, hi):
                        emit("content", slab=si)
                else:
                    for si in range(lo, hi):
                        emit("content", slab=si)
            else:
                for si in range(lo, hi):
                    emit("divider" if slabs[si]["kind"] == "divider" else "content", slab=si)
    else:
        for idx in range(len(slabs)):
            emit("content", slab=idx)
    for fb in file_blocks:
        fb["start_phys"] = slab_phys.get(fb["slab_lo"])

    def _arabic_at(slab_idx):
        cnt = 0
        for d in body_plan:
            if d["kind"] in ("divider", "blank", "content"):
                cnt += 1
            if d["slab"] == slab_idx:
                return cnt
        return cnt
    toc_entries = [(fb["name"], _arabic_at(fb["slab_lo"])) for fb in file_blocks]
    return body_plan, slab_phys, toc_entries


def _assemble(items, s, fast=False, gen_bookmarks_fn=None, tick=None, tmpdir=None):
    """표지+목차+(파일별 간지)+본문을 한 문서로 조립.
    번호: 표지=없음, 목차=로마자(i…), 본문=아라비아(1…, 간지·빈페이지는 카운트하되 숨김).
    양면+홀수시작: 각 문서(간지 포함)를 홀수 페이지에서 시작하도록 빈 페이지 삽입.
    반환: (final_doc, page_infos[{kind,cells}]).  final 에 책갈피 설정 완료."""
    s = merge_twoup_settings(s)
    nup = max(1, int(s.get("nup", 2) or 2))
    make_divider = bool(s.get("make_divider", False))
    doc_break = bool(s.get("doc_break", False))
    duplex = bool(s.get("duplex", False))
    per_file = make_divider or doc_break
    odd = bool(s.get("doc_start_odd", False)) and duplex and per_file
    make_cover = bool(s.get("make_cover", True))
    make_toc = bool(s.get("make_toc", True))

    def _tk(msg):
        if tick:
            tick(msg)

    opened = {}

    def _doc(it):
        k = id(it)
        if k not in opened:
            opened[k] = _open_item_cached(it)
        return opened[k][0]

    def _name(it):
        if isinstance(it, dict):
            return it.get("name") or Path(str(it.get("path", "항목"))).stem
        return Path(str(it)).stem

    try:
        # 1) 본문 슬랩(간지/내용) 구성
        slabs, file_blocks, page_to_slab = _merge_slabs(items, s, _item_pagecount)
        if not slabs:
            raise RuntimeError("배치할 내용이 없습니다.")

        # 본문 시트(앞장/간지/빈페이지)의 기준 크기 = 첫 내용 시트 크기
        def _first_ref():
            for sl in slabs:
                if sl["kind"] == "content":
                    for cell in sl["chunk"]:
                        if cell:
                            it, p = cell
                            return _doc(it)[p].rect
            sh, lo = _paper(s)
            return fitz.Rect(0, 0, sh, lo)
        fw, fh, _b0 = _grid_layout(_first_ref(), s)

        # 2) 앞장(표지·목차) 문서 생성 — 목차 쪽 수 확정용 1차
        cover_doc = None; toc_doc = None
        if make_cover:
            _tk("표지 생성 중…")
            if not fast and tmpdir is not None:
                cp = str(tmpdir / "cover.pdf"); cd = str(tmpdir / "cover.docx")
                if _fill_cover_docx(s.get("cover_template", ""), s.get("cover", {}), cd) \
                        and _docx_to_pdf(cd, cp):
                    try:
                        cover_doc = fitz.open(cp)
                    except Exception:
                        cover_doc = None
            if cover_doc is None:
                cover_doc = _fitz_cover_doc(s.get("cover", {}), fw, fh)
        cover_pages = cover_doc.page_count if cover_doc else 0

        names = [fb["name"] for fb in file_blocks]
        toc_pages = 0
        if make_toc:
            _tk("목차 생성 중…")
            toc_doc = _make_toc_doc([(n, "") for n in names], fw, fh, s, fast, tmpdir, 0)
            toc_pages = toc_doc.page_count
        front_pages = cover_pages + toc_pages

        # 3) 본문 번호/물리위치/빈페이지 계획
        body_plan, slab_phys, toc_entries = _merge_body_plan(slabs, file_blocks, s, front_pages)

        # 4) 목차 재생성(실제 번호) — 쪽 수 동일
        if make_toc:
            try:
                toc_doc.close()
            except Exception:
                pass
            toc_doc = _make_toc_doc(toc_entries, fw, fh, s, fast, tmpdir, 1)

        # 5) 조립
        final = fitz.open()
        page_infos = []
        if cover_doc:
            final.insert_pdf(cover_doc)
            page_infos += [{"kind": "cover", "cells": None}] * cover_pages
        if toc_doc:
            final.insert_pdf(toc_doc)
            page_infos += [{"kind": "toc", "cells": None}] * toc_doc.page_count

        # 간지 문서 캐시(파일별)
        div_docs = {}

        def _divider_for(fb):
            i = fb["file"]
            if i in div_docs:
                return div_docs[i]
            dd = None
            if not fast and tmpdir is not None:
                dp = str(tmpdir / f"div_{i}.pdf"); dx = str(tmpdir / f"div_{i}.docx")
                if _fill_divider_docx(s.get("divider_template", ""), fb["name"], dx) \
                        and _docx_to_pdf(dx, dp):
                    try:
                        dd = fitz.open(dp); _paint_bg_behind(dd, s.get("divider_bg", "#eef2f7"))
                    except Exception:
                        dd = None
            if dd is None:
                dd = _fitz_divider_doc(fb["name"], fw, fh, s.get("divider_bg", "#eef2f7"))
            div_docs[i] = dd
            return dd
        slab_file = {fb["slab_lo"]: fb for fb in file_blocks if slabs[fb["slab_lo"]]["kind"] == "divider"}

        for d in body_plan:
            if d["kind"] == "blank":
                bp = final.new_page(width=fw, height=fh); _draw_sheet_bg(bp, s)
                page_infos.append({"kind": "blank", "cells": None})
            elif d["kind"] == "divider":
                fb = slab_file.get(d["slab"])
                dd = _divider_for(fb)
                final.insert_pdf(dd, from_page=0, to_page=0)
                page_infos.append({"kind": "divider", "cells": None})
            else:  # content
                cells = _render_content_sheet(final, slabs[d["slab"]]["chunk"], s, d["phys"], _doc)
                page_infos.append({"kind": "content", "cells": cells})
            _tk("배치 중…")

        # 6) 번호 인쇄 — 표지 없음 / 목차 로마자 / 본문 아라비아(간지·빈페이지 숨김)
        idx = 0
        for _ in range(cover_pages):
            idx += 1
        for r in range(1, toc_pages + 1):
            pg = final[idx]; _draw_footer(pg, _roman(r), s, pg.rect.width); idx += 1
        for d in body_plan:
            pg = final[idx]
            if d["kind"] == "content":
                _draw_footer(pg, d["num"], s, pg.rect.width)
            idx += 1

        # 7) 책갈피
        bm = []
        for fb in file_blocks:
            bm.append((1, fb["name"], fb.get("start_phys") or (front_pages + 1)))
            it = fb["it"]
            emb = []
            try:
                emb = _doc(it).get_toc(simple=True) or []
            except Exception:
                emb = []
            if not emb and gen_bookmarks_fn is not None and isinstance(it, dict) \
                    and it.get("type") != "shots":
                try:
                    emb = [(int(lv) + 1, t, p) for (t, p, lv)
                           in gen_bookmarks_fn(str(it["path"]), _doc(it))]
                except Exception:
                    emb = []
            for lvl, title, pno in emb:
                slab = page_to_slab.get((id(it), max(1, int(pno)) - 1))
                if slab is not None and slab in slab_phys:
                    bm.append((max(2, int(lvl) + 1), title, slab_phys[slab]))
        toc_out = []; prev = 0
        for lvl, title, phys in bm:
            lv = 1 if not toc_out else min(max(1, int(lvl)), prev + 1)
            toc_out.append([lv, str(title) or "(제목 없음)", int(phys)])
            prev = lv
        try:
            if toc_out:
                final.set_toc(toc_out)
        except Exception:
            pass

        for dd in div_docs.values():
            try:
                dd.close()
            except Exception:
                pass
        if cover_doc:
            try:
                cover_doc.close()
            except Exception:
                pass
        if toc_doc:
            try:
                toc_doc.close()
            except Exception:
                pass
        # 260617-6: 맞쪽 인쇄 — 맨 앞에 여백(빈) 페이지 1장 추가(여백색). TOC 1쪽 보정.
        if bool(s.get("facing_first", False)) and final.page_count > 0:
            try:
                r0 = final[0].rect
                final.new_page(pno=0, width=r0.width, height=r0.height)
                _draw_sheet_bg(final[0], s)
                page_infos.insert(0, {"kind": "blank", "cells": None})
                t = final.get_toc(simple=True)
                if t:
                    final.set_toc([[lv, ti, pg + 1] for lv, ti, pg in t])
            except Exception:
                pass
        return final, page_infos
    finally:
        for dch, cached in opened.values():
            if not cached:
                try:
                    dch.close()
                except Exception:
                    pass
        _unpin_preview()


def _make_toc_doc(entries, w, h, s, fast, tmpdir, seq):
    if not fast and tmpdir is not None:
        tp = str(tmpdir / f"toc_{seq}.pdf"); tx = str(tmpdir / f"toc_{seq}.docx")
        if _fill_toc_docx(s.get("toc_template", ""), entries, tx) and _docx_to_pdf(tx, tp):
            try:
                return fitz.open(tp)
            except Exception:
                pass
    return _fitz_toc_doc(entries, w, h)


def _render_content_sheet(out_doc, chunk, s, phys_no, doc_fn):
    """내용 시트 1장 생성(방향·격자·채움·여백색·격자선). 셀 박스 목록 반환."""
    first = next((c for c in chunk if c is not None), None)
    if first is not None:
        ref = doc_fn(first[0])[first[1]].rect
    else:
        sh, lo = _paper(s); ref = fitz.Rect(0, 0, sh, lo)
    ow, oh, boxes = _grid_layout(ref, s, phys_no)
    pg = out_doc.new_page(width=ow, height=oh)
    _draw_sheet_bg(pg, s)
    for slot, cell in enumerate(chunk):
        if cell is None or slot >= len(boxes):
            continue
        it, pno = cell; d = doc_fn(it)
        if pno < d.page_count:
            _place_page(pg, d, pno, boxes[slot], s)
    _draw_grid_lines(pg, boxes, s)
    return [(b.x0, b.y0, b.x1, b.y1) for b in boxes]


def build_twoup(items, settings, out_path, gen_bookmarks_fn=None, log=None, progress=None):
    """items: [{type:'pdf'|'shots', path/paths, name}], settings: dict, out_path: str.
    gen_bookmarks_fn(path, src_doc) -> [(title, page1based, level)] (선택, 임베드 TOC 없을 때).
    progress(done, total, label) -> bool: False 반환 시 취소(MergeCancelled). 반환: out_path."""
    s = merge_twoup_settings(settings)
    tmpdir = Path(tempfile.mkdtemp(prefix="polypdf_2up_"))

    # 진행/취소: 총 작업량 ≈ 시트 수 + 표지/목차 + 저장
    nup = max(1, int(s.get("nup", 2) or 2))
    total_pages_est = 0
    for it in items:
        try:
            if isinstance(it, dict) and it.get("type") == "shots":
                total_pages_est += len(it.get("paths") or [])
            else:
                d0 = fitz.open(str(it["path"])); total_pages_est += d0.page_count; d0.close()
        except Exception:
            pass
    total_units = max(1, (total_pages_est + nup - 1) // nup + 4)
    done = [0]

    def _tick(label):
        done[0] += 1
        if progress is not None and progress(min(done[0], total_units), total_units, label) is False:
            raise MergeCancelled()

    final, _infos = _assemble(items, s, fast=False, gen_bookmarks_fn=gen_bookmarks_fn,
                              tick=_tick, tmpdir=tmpdir)
    _tick("최종 저장 중…")
    final.save(out_path, garbage=4, deflate=True)
    pages = final.page_count
    final.close()
    if log:
        log(f"다단 생성 완료: 총 {pages}쪽")
    return out_path

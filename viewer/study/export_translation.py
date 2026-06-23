"""번역 산출물 Word/PDF 조립 (P4·P4c) — 서지(APA) → 요약 → 전문 → 용어집 + 책갈피.

SOT: `PDF 번역·요약 작업 계획서.md` §9.
- docx: python-docx. 순서 = 서지(첫머리) → 2줄 → 요약 → (다음 장) 전문(섹션 제목=Heading) → 용어집.
- 그림/표는 **별도 목차가 아니라 본문에서 처음 언급되는 단락 뒤에 삽입**(원래 흐름 위치).
  그림 = 원문 이미지 + 번역 캡션, 표 = 원문 이미지 + **번역 표(실제 표 형식, 비전 재구성)** 병기.
  각 항목 캡션 = level-3 Heading → PDF 책갈피로도 탐색 가능.
- PDF: Word(win32com) ExportAsFixedFormat + 제목(Heading) 책갈피. Word 없으면 docx 만.
"""
from __future__ import annotations

import os
import re

_SEC_KW = ("서론", "서 론", "머리말", "개요", "배경", "재료", "방법", "실험", "분석",
           "결과", "고찰", "토의", "논의", "결론", "참고문헌", "감사의", "부록",
           "introduction", "method", "material", "result", "discussion",
           "conclusion", "reference")


def _is_heading(s: str) -> bool:
    s = (s or "").strip()
    if not s or len(s) > 40:
        return False
    if re.match(r"^\d+(?:\.\d+)*\.?\s+\S", s):       # "2. 재료 및 방법", "3.1 시험"
        return True
    if len(s) <= 24 and not s.endswith((".", "다.", "다", ":", ";", ",")):
        low = s.lower()
        if any(k in low for k in _SEC_KW):
            return True
    return False


def _set_korean_font(doc, name="맑은 고딕"):
    try:
        from docx.oxml.ns import qn
    except Exception:
        return
    for sname in ("Normal", "Title", "Heading 1", "Heading 2", "Heading 3"):
        try:
            st = doc.styles[sname]
            st.font.name = name
            rpr = st.element.get_or_add_rPr()
            rf = rpr.get_or_add_rFonts()
            rf.set(qn("w:eastAsia"), name)
        except Exception:
            pass


def _xml_clean(s: str) -> str:
    """docx/XML 비호환 제어문자 제거(PDF 추출 텍스트는 NULL/제어문자 포함 가능)."""
    return "".join(ch for ch in (s or "")
                   if ch in "\t\n\r" or ord(ch) >= 0x20)


_CAP_LABEL_RE = re.compile(r"^\s*(?:Tables?|Figures?|Figs?|표|그림)\s*\d+\s*[.:]?\s*", re.I)


def _compose_caption(num, kind, caption_ko, caption_en=""):
    """캡션 한글 단일화 — 영문 라벨/중복 제거 후 `표 N. 제목` / `그림 N. 제목`."""
    title = (caption_ko or "").strip()
    title = _CAP_LABEL_RE.sub("", title).strip()         # 'Table 1.'·'표 1.' 라벨 제거
    if not title and caption_en:
        title = _CAP_LABEL_RE.sub("", caption_en.strip()).strip()
    label = "표" if kind == "tab" else "그림"
    head = f"{label} {num}." if num else label
    return (head + " " + title).strip() if title else head


def _strip_grid_caption(rows, num):
    """비전 표 그리드 첫 행이 캡션('표/Table N …')이면 제거(중복 방지)."""
    if not rows:
        return rows
    first = " ".join(c for c in rows[0] if c).strip()
    nonempty = [c for c in rows[0] if (c or "").strip()]
    if len(nonempty) <= 2 and re.search(r"(?:Tables?|표|Figures?|그림)\s*\d+", first, re.I):
        return rows[1:]
    return rows


def _strip_md(s: str) -> str:
    """요약 등 본문 줄의 마크다운 마크업을 평문으로 정리(`**`,`##`,`|`,`---`)."""
    s = (s or "")
    if set(s.strip()) <= set("-—–=* ") and s.strip():
        return ""                                         # 구분선
    s = re.sub(r"\*\*(.+?)\*\*", r"\1", s)               # 굵게
    s = re.sub(r"`([^`]+)`", r"\1", s)                    # 인라인 코드
    s = re.sub(r"^\s*#{1,6}\s*", "", s)                   # 헤딩 마크
    s = re.sub(r"^\s*[-*]\s+", "• ", s)                   # 불릿
    s = s.replace("|", " ")
    return re.sub(r"[ \t]{2,}", " ", s).strip()


def _anchor_index(paras, num, kind):
    """본문 단락 중 그림/표 N 을 처음 언급하는 단락 인덱스(없으면 None).

    kind='fig' → 'Figure/Fig/그림 N', 'tab' → 'Table/표 N'(뒤에 숫자 경계)."""
    if not num:
        return None
    if kind == "fig":
        rx = re.compile(r"(?:Figures?|Figs?\.?|그림)\s*0*%d(?!\d)" % num, re.I)
    else:
        rx = re.compile(r"(?:Tables?|표)\s*0*%d(?!\d)" % num, re.I)
    for i, p in enumerate(paras):
        if rx.search(p):
            return i
    return None


def _render_table_grid(doc, rows):
    """번역된 표 그리드(rows=list[list[str]])를 실제 Word 표로 렌더."""
    rows = [r for r in (rows or []) if any((c or "").strip() for c in r)]
    if not rows:
        return False
    ncol = max(len(r) for r in rows)
    if ncol < 1:
        return False
    t = doc.add_table(rows=0, cols=ncol)
    try:
        t.style = "Table Grid"
    except Exception:
        pass
    for r in rows:
        cells = t.add_row().cells
        for j in range(ncol):
            cells[j].text = _xml_clean(r[j]) if j < len(r) else ""
    return True


def _render_asset(doc, kind, item):
    """그림/표 1개를 본문 흐름 안에 삽입.

    그림 = 이미지 → 그 아래 한글 캡션(=책갈피). 표 = 한글 캡션 → 바로 아래 번역 표(docx).
    영문 캡션·[원문]/[번역] 라벨·표 원문 이미지는 두지 않는다(번역 표 실패 시에만 원문 폴백)."""
    from docx.shared import Inches
    num = item.get("num")
    head = _xml_clean(_compose_caption(num, kind, item.get("caption_ko"),
                                       item.get("caption")))[:90]
    imgs = [im for im in (item.get("images") or [item.get("image")])
            if im and os.path.exists(im)]
    if kind == "fig":
        # 그림(원문 이미지들, 연속 그림이면 모두) → 아래 한글 캡션
        for im in imgs:
            try:
                doc.add_picture(im, width=Inches(5.7))
            except Exception:
                pass
        doc.add_heading(head, level=3)
        doc.add_paragraph("")                # 그림 제목 아래 본문이 바로 붙지 않도록 빈 줄
        return
    # 표: 한글 캡션 → 바로 아래 번역 표(그리드, 연속 표는 병합된 행)
    doc.add_heading(head, level=3)
    rows = _strip_grid_caption(item.get("rows_ko") or [], num)
    if rows:
        _render_table_grid(doc, rows)
    else:
        for im in imgs:                      # 번역 표 실패 시 원문 이미지(들) 폴백
            try:
                doc.add_picture(im, width=Inches(6.2))
            except Exception:
                pass
    doc.add_paragraph("")                    # 표 아래 본문이 바로 붙지 않도록 빈 줄


_EQ_TOK_RE = re.compile(r"【수식(\d+)】")


def _render_para_with_eq(doc, para, eq_map):
    """단락에 【수식N】 토큰이 있으면 텍스트/수식 이미지를 순서대로 렌더(번역 안 한 수식 이미지).

    수식 이미지는 **원본 폭(pt→inch)** 에 맞춰 크기 지정(과대 확대 방지, 최대 6인치)."""
    from docx.shared import Inches
    segs = _EQ_TOK_RE.split(para)        # [text, num, text, num, …]
    for i, seg in enumerate(segs):
        if i % 2 == 1:
            ent = eq_map.get(int(seg))
            im = ent[0] if isinstance(ent, tuple) else ent
            wpt = ent[1] if isinstance(ent, tuple) else 0
            if im and os.path.exists(im):
                w = min(6.0, max(1.0, (wpt or 260) / 72.0))
                try:
                    doc.add_picture(im, width=Inches(w))
                except Exception:
                    pass
        else:
            s = (_strip_md(seg) or "").strip()
            if s:
                doc.add_paragraph(s)


def _emit_translation_body(doc, translation, figures, tables, equations=None):
    """전문 번역 단락을 흐름대로 출력하되, 그림/표는 처음 언급 단락 뒤, 수식은 토큰 위치에 삽입."""
    eq_map = {e["id"]: (e.get("image"), e.get("width_pt", 0))
              for e in (equations or []) if e.get("image")}
    paras = [p.strip() for p in (translation or "").split("\n\n") if p.strip()]
    assets = [("fig", f) for f in (figures or [])] + [("tab", t) for t in (tables or [])]
    by_idx = {}
    leftover = []
    for kind, item in assets:
        idx = _anchor_index(paras, item.get("num"), kind)
        if idx is None:
            leftover.append((kind, item))
        else:
            by_idx.setdefault(idx, []).append((kind, item))
    for i, para in enumerate(paras):
        if eq_map and _EQ_TOK_RE.search(para):
            _render_para_with_eq(doc, para, eq_map)       # 수식 이미지(번역 안 함) 인라인
        else:
            para = _strip_md(para) or para
            if _is_heading(para):
                doc.add_heading(para, level=2)
            else:
                doc.add_paragraph(para)
        for kind, item in by_idx.get(i, []):
            _render_asset(doc, kind, item)
    # 본문에서 참조되지 않은 그림/표는 끝에 모아 배치
    if leftover:
        doc.add_heading("기타 그림·표", level=2)
        for kind, item in leftover:
            _render_asset(doc, kind, item)


def build_docx(out_path, *, title="", citation="", summary="", translation="",
               glossary=None, figures=None, tables=None, equations=None):
    """구조화 docx 저장(서지 → 요약 → 전문[그림/표/수식 인라인 삽입] → 용어집)."""
    from docx import Document
    doc = Document()
    _set_korean_font(doc)

    # 1) 첫머리: 서지(APA) — 본 논문 자체
    if (citation or "").strip():
        p = doc.add_paragraph()
        run = p.add_run(citation.strip())
        run.bold = True
    elif title:
        doc.add_paragraph(str(title)).runs and None

    # 2) 2줄 띄고 → 요약(같은 페이지)
    doc.add_paragraph("")
    doc.add_paragraph("")
    if (summary or "").strip():
        doc.add_heading("요약", level=1)
        for ln in summary.splitlines():
            ln = _strip_md(ln.rstrip())
            if ln:
                doc.add_paragraph(ln)

    # 3) 요약이 끝나면 다음 장부터 전문 번역(섹션 제목 = Heading)
    #    그림/표는 별도 목차가 아니라 본문에서 처음 언급되는 위치에 삽입(번역 표 = 실제 표 형식)
    doc.add_page_break()
    doc.add_heading("전문 번역", level=1)
    _emit_translation_body(doc, translation, figures, tables, equations)

    # 4) 용어집 부록
    gl = [g for g in (glossary or []) if g.get("en") and g.get("ko")]
    if gl:
        doc.add_page_break()
        doc.add_heading("용어집", level=1)
        t = doc.add_table(rows=1, cols=3)
        try:
            t.style = "Table Grid"
        except Exception:
            pass
        hc = t.rows[0].cells
        hc[0].text, hc[1].text, hc[2].text = "원어(EN)", "번역(KO)", "출처"
        for g in gl[:200]:
            c = t.add_row().cells
            c[0].text = g.get("en", "")
            c[1].text = g.get("ko", "")
            c[2].text = g.get("source", "")

    doc.save(out_path)
    return out_path


def docx_to_pdf(docx_path, pdf_path):
    """(성공, 메시지). Word(win32com)로 PDF 변환 + 제목(Heading) 책갈피. 스레드 안전(CoInitialize)."""
    pythoncom = None
    word = None
    try:
        import pythoncom
        import win32com.client
        pythoncom.CoInitialize()
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        d = word.Documents.Open(os.path.abspath(docx_path))
        # 17=wdExportFormatPDF, CreateBookmarks=1=wdExportCreateHeadingBookmarks
        d.ExportAsFixedFormat(OutputFileName=os.path.abspath(pdf_path),
                              ExportFormat=17, CreateBookmarks=1)
        d.Close(False)
        return True, ""
    except Exception as e:
        return False, f"{type(e).__name__}: {str(e)[:100]}"
    finally:
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        try:
            if pythoncom is not None:
                pythoncom.CoUninitialize()
        except Exception:
            pass


def _safe_name(name: str) -> str:
    return (re.sub(r'[\\/:*?"<>|]+', "_", name or "").strip() or "번역")[:120]


def _unique_pair(folder, base):
    """({base}.docx, {base}.pdf) 가 비어 있으면 그대로, 충돌하면 '{base} (k)' 로 회피.

    docx·pdf 는 동일 접미를 쓰도록 둘 중 하나라도 있으면 다음 번호로(Windows 방식).
    이전 번역본 보존 + PDF 가 열려 잠겨 있어도 새 파일로 안전 저장."""
    d = os.path.join(folder, base + ".docx")
    p = os.path.join(folder, base + ".pdf")
    if not os.path.exists(d) and not os.path.exists(p):
        return d, p
    for k in range(1, 1000):
        d = os.path.join(folder, f"{base} ({k}).docx")
        p = os.path.join(folder, f"{base} ({k}).pdf")
        if not os.path.exists(d) and not os.path.exists(p):
            return d, p
    return d, p          # 폴백(거의 도달 안 함)


def glossary_sidecar_path(folder, name) -> str:
    """그 PDF 의 번역 용어집 사이드카 경로 `{이름}_용어집.json`(번역 방식 무관 교정용)."""
    return os.path.join(folder, f"{_safe_name(name)}_용어집.json")


def save_glossary_sidecar(folder, name, glossary):
    """번역에 사용된 용어집을 사이드카 JSON 으로 저장(나중에 '용어집 교정'에서 불러옴)."""
    import json
    gl = [{"en": g.get("en", ""), "ko": g.get("ko", ""), "source": g.get("source", "")}
          for g in (glossary or []) if g.get("en") and g.get("ko")]
    try:
        with open(glossary_sidecar_path(folder, name), "w", encoding="utf-8") as f:
            json.dump({"name": name, "glossary": gl}, f, ensure_ascii=False, indent=0)
    except Exception:
        pass


def load_glossary_sidecar(folder, name):
    """사이드카 용어집 로드 → [{en, ko, source}]. 없으면 []."""
    import json
    try:
        with open(glossary_sidecar_path(folder, name), encoding="utf-8") as f:
            return (json.load(f) or {}).get("glossary") or []
    except Exception:
        return []


def resolve_glossary_sidecar(path):
    """원본 PDF 또는 번역본(`{이름}_번역[ (k)].pdf`) 경로 → 존재하는 사이드카 경로('' 없으면).

    번역된 자료에서 우클릭해도 용어집 교정을 열 수 있도록 '_번역[ (k)]' 접미를 벗겨 원본명으로도 탐색."""
    p = os.path
    folder, fn = p.dirname(str(path)), p.splitext(p.basename(str(path)))[0]
    cands = [fn]
    m = re.match(r"^(.*)_번역(?: \(\d+\))?$", fn)
    if m:
        cands.append(m.group(1))
    for base in cands:
        sc = glossary_sidecar_path(folder, base)
        if p.exists(sc):
            return sc
    return ""


def save_translation_doc(folder, name, *, citation="", summary="", translation="",
                         glossary=None, figures=None, tables=None, equations=None):
    """(docx_path, pdf_path|'', [진단]). 논문 폴더에 _번역.docx/.pdf 저장(충돌 시 (k) 회피)."""
    safe = _safe_name(name)
    docx_path, pdf_path = _unique_pair(folder, f"{safe}_번역")
    build_docx(docx_path, title=name, citation=citation, summary=summary,
               translation=translation, glossary=glossary,
               figures=figures, tables=tables, equations=equations)
    save_glossary_sidecar(folder, name, glossary)      # 용어집 교정용 사이드카
    ok, msg = docx_to_pdf(docx_path, pdf_path)
    dbg = ["docx 저장" + (" + PDF(책갈피)" if ok else f" (PDF 변환 실패: {msg})")]
    return docx_path, (pdf_path if ok else ""), dbg

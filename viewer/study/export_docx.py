"""단어장 내용을 Word(.docx) 로 저장. python-docx 사용. 정렬·표시옵션 반영."""
from __future__ import annotations

import re
from pathlib import Path

_HANGUL = re.compile(r"[가-힣]")


def _sort_rows(rows, sort):
    if sort == "가나다·ABC":
        return sorted(rows, key=lambda r: r["lemma"])
    if sort == "빈도순":
        return sorted(rows, key=lambda r: (-r.get("count", 1), r["lemma"]))
    return sorted(rows, key=lambda r: r.get("pos", 0))   # 문장 순서


def export_study_docx(store, file_key, out_path, *,
                      title: str = "단어장", levels=None, user_overrides=None,
                      sort: str = "문장 순서",
                      show_ko: bool = True, show_en: bool = True,
                      show_ex: bool = True,
                      progress=None) -> Path:
    """파일의 전체 단어(페이지별)를 .docx 로 저장. 정렬·표시옵션 반영. 반환: 출력 경로."""
    from docx import Document
    from docx.shared import Pt, RGBColor

    out_path = Path(out_path)
    doc = Document()
    doc.add_heading(title, level=0)

    pages = [r["page"] for r in store.conn.execute(
        "SELECT DISTINCT page FROM vocab_page WHERE file_key=? ORDER BY page",
        (file_key,))]
    color = {"초급": RGBColor(0x2e, 0xa0, 0x44), "중급": RGBColor(0xd0, 0x9b, 0x00),
             "고급": RGBColor(0xc0, 0x39, 0x2b), "미정": RGBColor(0x88, 0x88, 0x88)}
    total = 0
    for pi, pno in enumerate(pages):
        if progress:
            progress(pi, len(pages))
        rows = store.get_page_study(file_key, pno, levels=levels,
                                    user_overrides=user_overrides)
        rows = _sort_rows(rows, sort)
        if not rows:
            continue
        doc.add_heading(f"p.{pno + 1}", level=2)
        for r in rows:
            lv = r.get("level", "미정")
            p = doc.add_paragraph()
            rb = p.add_run("● ")
            rb.font.color.rgb = color.get(lv, color["미정"])
            rw = p.add_run(r["lemma"])
            rw.bold = True
            rw.font.size = Pt(12)
            if r.get("count", 1) > 1:
                p.add_run(f"  ×{r['count']}").font.size = Pt(9)
            defs = r.get("definitions") or []
            ko_defs = [d for d in defs if _HANGUL.search(d["definition"])]
            en_defs = [d for d in defs if not _HANGUL.search(d["definition"])]
            if show_ko:
                for d in ko_defs:
                    doc.add_paragraph(d["definition"], style="List Bullet")
            if show_en:
                for d in en_defs:
                    doc.add_paragraph(d["definition"], style="List Bullet")
            if show_ex:
                for e in (r.get("examples") or []):
                    ep = doc.add_paragraph(style="List Bullet 2")
                    ep.add_run("예) " + e["example"]).italic = True
            total += 1

    doc.add_paragraph()
    doc.add_paragraph(f"총 {total} 단어 · {len(pages)} 페이지 · 정렬: {sort}").italic = True
    doc.save(str(out_path))
    return out_path
